from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from openpyxl import load_workbook

BIN_PATTERN = re.compile(r"\b(\d{12})\b")
DASH_SPLIT = re.compile(r"\s*[-–—]\s*")

NAME_HEADER_HINTS = ("name", "название", "наименование", "компания")
IIN_HEADER_HINTS = ("iinbin", "iin", "bin", "иин", "бин", "рнн")


def _normalize_header(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip().lower())


def _header_matches(header: str, hints: tuple[str, ...]) -> bool:
    normalized = _normalize_header(header)
    return any(h in normalized for h in hints)


def _looks_like_header_row(cells: list[str]) -> bool:
    """True only for short label rows without embedded BINs (not company names)."""
    if not cells or any(BIN_PATTERN.search(cell) for cell in cells):
        return False
    return any(
        _header_matches(cell, NAME_HEADER_HINTS) or _header_matches(cell, IIN_HEADER_HINTS)
        for cell in cells
        if len(_normalize_header(cell)) <= 40
    )


def _find_column_index(headers: list[str], hints: tuple[str, ...], default: int) -> int:
    for index, header in enumerate(headers):
        if _header_matches(header, hints):
            return index
    return default


def _digits_only(value: str) -> str:
    return re.sub(r"\D", "", value or "")


def _name_from_freeform_line(text: str, iin_bin: str) -> str:
    without_bin = BIN_PATTERN.sub("", text).strip()
    without_bin = re.sub(r"\s*[-–—]\s*$", "", without_bin).strip()
    parts = [part.strip() for part in DASH_SPLIT.split(without_bin) if part.strip()]
    if len(parts) >= 2:
        return " - ".join(parts[1:])
    return without_bin or text.strip()


def _row_from_cells(cells: list[str]) -> dict[str, str] | None:
    if not cells:
        return None

    if len(cells) == 1:
        text = cells[0].strip()
        if not text:
            return None
        match = BIN_PATTERN.search(text)
        if not match:
            return None
        iin_bin = match.group(1)
        name = _name_from_freeform_line(text, iin_bin)
        return {"name": name, "iinBin": iin_bin}

    if _looks_like_header_row(cells):
        return None

    name_idx = _find_column_index(cells, NAME_HEADER_HINTS, 0)
    iin_idx = _find_column_index(cells, IIN_HEADER_HINTS, 1 if len(cells) > 1 else 0)
    name = cells[name_idx].strip() if name_idx < len(cells) else ""
    iin_raw = cells[iin_idx].strip() if iin_idx < len(cells) else ""
    iin_bin = _digits_only(iin_raw)
    if not iin_bin and name:
        match = BIN_PATTERN.search(name)
        if match:
            iin_bin = match.group(1)
            name = _name_from_freeform_line(name, iin_bin)
    if name and iin_bin:
        return {"name": name, "iinBin": iin_bin}
    return None


def parse_excel_bytes(content: bytes) -> list[dict[str, str]]:
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sheet = workbook.active
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return []

    header = [str(cell or "").strip() for cell in rows[0]]
    name_idx = _find_column_index(header, NAME_HEADER_HINTS, 0)
    iin_idx = _find_column_index(header, IIN_HEADER_HINTS, 1 if len(header) > 1 else 0)

    items: list[dict[str, str]] = []
    for row in rows[1:]:
        if not row:
            continue
        cells = [str(cell or "").strip() for cell in row]
        name = cells[name_idx] if name_idx < len(cells) else ""
        iin_raw = cells[iin_idx] if iin_idx < len(cells) else ""
        iin_bin = _digits_only(iin_raw)
        if name and iin_bin:
            items.append({"name": name, "iinBin": iin_bin})
    return items


def _table_rows(document: Document) -> list[list[str]]:
    table_rows: list[list[str]] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_rows.append(cells)
    return table_rows


def _paragraph_lines(document: Document) -> list[str]:
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    return lines


def _rows_from_table(table_rows: list[list[str]]) -> list[dict[str, str]]:
    if not table_rows:
        return []

    header = table_rows[0]
    if _looks_like_header_row(header):
        name_idx = _find_column_index(header, NAME_HEADER_HINTS, 0)
        iin_idx = _find_column_index(header, IIN_HEADER_HINTS, 1 if len(header) > 1 else 0)
        data_rows = table_rows[1:]
        items: list[dict[str, str]] = []
        for cells in data_rows:
            if not any(cells):
                continue
            name = cells[name_idx].strip() if name_idx < len(cells) else ""
            iin_raw = cells[iin_idx].strip() if iin_idx < len(cells) else ""
            iin_bin = _digits_only(iin_raw)
            if name and iin_bin:
                items.append({"name": name, "iinBin": iin_bin})
        if items:
            return items

    items = []
    for cells in table_rows:
        parsed = _row_from_cells(cells)
        if parsed:
            items.append(parsed)
    return items


def parse_docx_bytes(content: bytes) -> list[dict[str, str]]:
    document = Document(io.BytesIO(content))
    items: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()

    table_rows = _table_rows(document)
    if table_rows:
        for parsed in _rows_from_table(table_rows):
            key = (parsed["name"], parsed["iinBin"])
            if key not in seen:
                seen.add(key)
                items.append(parsed)
        if items:
            return items

    for line in _paragraph_lines(document):
        parsed = _row_from_cells([line])
        if parsed:
            key = (parsed["name"], parsed["iinBin"])
        else:
            key = (line, "")
            parsed = {"name": line, "iinBin": ""}
        if key not in seen:
            seen.add(key)
            items.append(parsed)

    return items


def parse_import_file(filename: str, content: bytes) -> list[dict[str, str]]:
    lower = filename.lower()
    if lower.endswith(".docx"):
        return parse_docx_bytes(content)
    if lower.endswith((".xlsx", ".xls")):
        return parse_excel_bytes(content)
    if content[:2] == b"PK":
        return parse_docx_bytes(content)
    return parse_excel_bytes(content)


def validate_import_row(name: str, iin_bin: str) -> dict[str, Any]:
    digits = _digits_only(iin_bin)
    extra_data: dict[str, str] = {}
    valid = True
    error: str | None = None

    if not name.strip():
        valid = False
        error = "Отсутствует название"
    elif not digits:
        valid = False
        error = "Отсутствует ИИН/БИН"
    elif len(digits) != 12:
        valid = False
        error = "ИИН/БИН должен содержать 12 цифр"

    return {
        "name": name.strip(),
        "iinBin": digits,
        "extraData": extra_data,
        "valid": valid,
        "error": error,
    }


def preview_import_rows(items: list[dict[str, str]]) -> list[dict[str, Any]]:
    return [validate_import_row(item["name"], item["iinBin"]) for item in items]
